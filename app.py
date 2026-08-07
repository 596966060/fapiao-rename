#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
发票批量重命名工具 - 完整版
支持：上传 → 识别 → 重命名 → 下载重命名后的发票文件
"""

from flask import Flask, render_template, request, jsonify, send_file
import os
import sys
import tempfile
import shutil
import zipfile
from pathlib import Path
import re
import io
from datetime import datetime
import threading
import openpyxl
from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                              GradientFill)
from openpyxl.utils import get_column_letter

# PDF 和图像处理
from PIL import Image
import numpy as np
import cv2

# OCR
import easyocr

# Windows 编码
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ======================== 发票提取 ========================

class InvoiceExtractor:
    """发票 OCR 提取"""

    def __init__(self, reader):
        self.reader = reader

    def _pdf_to_image(self, pdf_path: str) -> np.ndarray:
        """PDF 转图片"""
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=150)
            if not images:
                raise Exception("PDF 无法转换")
            image = images[0]
            return cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        except ImportError:
            raise Exception("缺少 pdf2image 库")
        except Exception as e:
            raise Exception(f"PDF 转图片失败: {e}")

    def _image_file_to_array(self, image_path: str) -> np.ndarray:
        """图片转数组"""
        try:
            image = Image.open(image_path).convert('RGB')
            return cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        except Exception as e:
            raise Exception(f"图片读取失败: {e}")

    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """图像预处理 - 仅做对比度增强，不做二值化（会破坏OCR）"""
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            # 自适应对比度增强
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)
            # 返回增强后的灰度图转BGR
            return cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)
        except:
            return image

    def _extract_from_filename(self, stem: str) -> dict:
        """从文件名中提取发票信息作为补充（当OCR识别不全时使用）
        常见格式: dzfp_263220000004482288706_上海松鼠创科技术有限责任公司_20260604112321
        """
        result = {}
        parts = stem.split('_')

        for part in parts:
            # 提取发票号：15-25位纯数字
            if re.match(r'^\d{15,25}$', part):
                result.setdefault('invoice_number', part)
            # 提取日期：8位紧凑格式 20260604，严格月日校验
            elif re.match(r'^20\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d*$', part):
                date_str = part[:8]
                y, m, d = int(date_str[:4]), int(date_str[4:6]), int(date_str[6:8])
                if 1 <= m <= 12 and 1 <= d <= 31:
                    result.setdefault('date', f"{y:04d}-{m:02d}-{d:02d}")
            # 提取公司名：包含中文且带有企业词尾
            elif re.search(r'[\u4e00-\u9fa5]', part) and len(part) >= 4:
                result.setdefault('buyer', part)

        return result

    def extract(self, file_path: str) -> dict:
        """提取发票信息 - 支持多个OCR引擎降级"""
        try:
            file_path = Path(file_path)
            ext = file_path.suffix.lower()

            if ext == ".pdf":
                image_array = self._pdf_to_image(str(file_path))
            else:
                image_array = self._image_file_to_array(str(file_path))

            if image_array is None or image_array.size == 0:
                raise Exception("无法读取图像")

            image_array = self._preprocess_image(image_array)

            # 尝试 EasyOCR
            results = self.reader.readtext(image_array, detail=0)
            if not results:
                raise Exception("OCR 无法识别")

            text = '\n'.join(results)

            fields = self._extract_fields(text)

            # 如果关键字段缺失，尝试从文件名补充
            filename_fields = self._extract_from_filename(file_path.stem)
            for key in ('date', 'invoice_number', 'buyer', 'supplier'):
                if not fields.get(key) and filename_fields.get(key):
                    fields[key] = filename_fields[key]

            # 保留原始 OCR 文本供前端调试展示
            fields['_raw_text'] = text
            return fields
        except Exception as e:
            raise Exception(f"提取失败: {e}")

    # ------------------------------------------------------------------ #
    # 文本预处理：修复 OCR 两栏布局拆行 & 统一各发票类型的字段标签
    # ------------------------------------------------------------------ #
    @staticmethod
    def _normalize_text(text: str) -> str:
        """
        修复常见 OCR 问题：
        1. 两栏发票中购/销方标签被单独拆成一行，下一行才是"名称:"
           "购\n名称:XX" → "购买方名称:XX"
        2. "销售方信息" 作为大标题，下面跟"名称:" 字段
        3. 机动车/航空/其他发票格式兼容
        4. OCR 误读修复：购→构/沟，销→消/晓，名→1/l 等
        5. 支持"客户名称""收款方"等替代标签统一为购买方/销售方
        """
        lines = text.split('\n')
        out = []
        i = 0
        while i < len(lines):
            raw = lines[i]
            stripped = raw.strip()

            # ---- OCR 误读修复（逐行预处理）----
            # 购买方相关：构买方/沟买方/构方/沟方/构买万 → 购买方
            stripped = re.sub(r'^(构|沟)(?:买?方?)(?:信息)?$', '购买方', stripped)
            stripped = re.sub(r'^(构|沟)(?:买?方?)\s*(?:信息)?$', '购买方', stripped)
            # 销售方相关：消售方/晓售方/消方/消售万 → 销售方
            stripped = re.sub(r'^(消|晓)(?:售?方?)(?:信息)?$', '销售方', stripped)
            stripped = re.sub(r'^(消|晓)(?:售?方?)\s*(?:信息)?$', '销售方', stripped)
            # 买方/卖方（住宿发票常用）
            stripped = re.sub(r'^(构|沟)(?:买)(?:万|放)?(?:信息)?$', '买方', stripped)
            stripped = re.sub(r'^(消|晓)(?:卖)(?:万|放)?(?:信息)?$', '卖方', stripped)
            # 名称: → OCR 误读 "1称:" / "l称:" / "1秄:" / "名称∶"
            stripped = re.sub(r'^([1l])\s*[称秄][：:﹕∶]', '名称:', stripped)
            # 销售方信息（大字标题）→ 统一
            stripped = re.sub(r'^销\s*售\s*方\s*信\s*息', '销售方信息', stripped)
            stripped = re.sub(r'^购\s*买\s*方\s*信\s*息', '购买方信息', stripped)
            # OCR 把"售"读成"害/室/富"，把"方"读成"万/放/仿"
            stripped = re.sub(r'^销[害室富][万放仿]?(?:信息)?$', '销售方', stripped)
            stripped = re.sub(r'^购[害室富]?[万放仿]?(?:信息)?$', '购买方', stripped)

            # ---- 替代标签统一 ----
            # "客户名称:" / "客户名称" → 购买方名称:
            stripped = re.sub(r'客户名称[：:]', '购买方名称:', stripped)
            stripped = re.sub(r'客户名称$', '购买方名称:', stripped)
            # "付款方:" → 购买方名称:
            stripped = re.sub(r'付款方[：:]', '购买方名称:', stripped)
            stripped = re.sub(r'付款方$', '购买方名称:', stripped)
            # "收款方:" → 销售方名称:
            stripped = re.sub(r'收款方[：:]', '销售方名称:', stripped)
            stripped = re.sub(r'收款方$', '销售方名称:', stripped)
            # "供应方:" → 销售方名称:
            stripped = re.sub(r'供应方[：:]', '销售方名称:', stripped)
            stripped = re.sub(r'供应方$', '销售方名称:', stripped)
            # "卖方名称:" → 销售方名称:
            stripped = re.sub(r'卖方[名1l][称秄][：:]', '销售方名称:', stripped)
            stripped = re.sub(r'卖方$', '销售方名称:', stripped)
            # "买方名称:" → 购买方名称:
            stripped = re.sub(r'买方[名1l][称秄][：:]', '购买方名称:', stripped)
            stripped = re.sub(r'买方$', '购买方名称:', stripped)

            # ---- 跨行标签重建 ----
            # 向前看：下一行是否以"名称:" 开头（含OCR常见误读）
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                nxt_is_name = bool(re.match(r'[名1l]\s*[称秄][：:﹕]', nxt))
                if nxt_is_name:
                    # 孤立的购方标签（含OCR误读变体）
                    if re.fullmatch(r'(?:购|构|沟)(?:买?方?|方)?(?:信息)?', stripped):
                        out.append('购买方' + nxt)
                        i += 2
                        continue
                    # 孤立的销方标签（含OCR误读变体）
                    if re.fullmatch(r'(?:销|消|晓)(?:售?方?|方)?(?:信息)?', stripped):
                        out.append('销售方' + nxt)
                        i += 2
                        continue
                    # 孤立的"买"标签（住宿发票常用"买方"）
                    if re.fullmatch(r'(?:买|构买)', stripped):
                        out.append('购买方' + nxt)
                        i += 2
                        continue
                    # 孤立的"卖/销"标签
                    if re.fullmatch(r'(?:卖|消|晓|销)', stripped):
                        out.append('销售方' + nxt)
                        i += 2
                        continue
                    # 孤立的"客户"标签
                    if re.fullmatch(r'客户', stripped):
                        out.append('购买方' + nxt)
                        i += 2
                        continue
                    # 孤立的"收款"标签
                    if re.fullmatch(r'收款', stripped):
                        out.append('销售方' + nxt)
                        i += 2
                        continue

            # ---- 机动车发票标签统一 ----
            stripped = re.sub(r'销货单位名称', '销售方名称', stripped)
            stripped = re.sub(r'购货单位名称', '购买方名称', stripped)
            stripped = re.sub(r'销货单位[：:]', '销售方名称：', stripped)
            stripped = re.sub(r'购货单位[：:]', '购买方名称：', stripped)

            out.append(raw[:raw.find(lines[i].lstrip())] + stripped if stripped != raw.strip() else raw)
            i += 1
        return '\n'.join(out)

    def _extract_special_invoice(self, text: str, result: dict) -> bool:
        """
        识别并提取特殊发票类型，返回 True 表示已处理（后续可跳过通用逻辑）。
        支持：机动车销售统一发票、航空运输电子客票行程单、出租车发票、定额发票
        """
        # ---- 机动车销售统一发票 ----
        if re.search(r'机动车销售统一发票|机动车(?:出售|发票)', text):
            result['invoice_type'] = '机动车发票'
            # 购货单位 / 销货单位（已在_normalize_text中替换为标准标签，无需重复处理）
            return False   # 让通用逻辑继续跑购买方/销售方

        # ---- 航空运输电子客票行程单 ----
        if re.search(r'航空运输电子客票|行程单|旅客姓名|电子客票', text):
            result['invoice_type'] = '航空行程单'
            # 旅客姓名 → buyer
            m = re.search(r'旅客姓名[：:\s]+([^\n\s]{2,10})', text)
            if m:
                result['buyer'] = m.group(1).strip()
            # 出发地-目的地 → supplier 位置存路线（用于文件名）
            dep = re.search(r'出发地[：:\s]+([^\n\s]{2,8})', text)
            arr = re.search(r'目的地[：:\s]+([^\n\s]{2,8})', text)
            if dep and arr:
                result['supplier'] = f"{dep.group(1).strip()}-{arr.group(1).strip()}"
            elif dep:
                result['supplier'] = dep.group(1).strip()
            # 票价
            m = re.search(r'(?:票价|合计)[：:\s]*[¥￥]?\s*(\d+(?:\.\d{1,2})?)', text)
            if m:
                result['amount'] = f"{float(m.group(1)):.2f}"
            return True

        # ---- 出租车发票 ----
        if re.search(r'出租汽车|出租车发票|计价器', text):
            result['invoice_type'] = '出租车发票'
            # 出租车发票通常有"客户名称""付款方"作为购买方，"收款方""供应方"作为销售方
            # 优先从 _normalize_text 处理后的文本中提取（标签已统一）
            buyer_m = re.search(r'(?:购买方|客户|付款方)[：:\s]*([\u4e00-\u9fa5]{2,30}(?:公司|集团|有限|企业|中心)?)', text)
            if buyer_m:
                result.setdefault('buyer', buyer_m.group(1).strip())
            supplier_m = re.search(r'(?:销售方|收款方|供应方|运营公司)[：:\s]*([\u4e00-\u9fa5]{2,30}(?:公司|集团|有限|企业|中心)?)', text)
            if supplier_m:
                result.setdefault('supplier', supplier_m.group(1).strip())
            # 金额
            m = re.search(r'(?:金额|合计|票价)[：:\s]*[¥￥]?\s*(\d+(?:\.\d{1,2})?)', text)
            if m:
                result['amount'] = f"{float(m.group(1)):.2f}"
            # 日期
            if not result.get('date'):
                dm = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})', text)
                if dm:
                    result['date'] = f"{dm.group(1)}-{int(dm.group(2)):02d}-{int(dm.group(3)):02d}"
            return True

        # ---- 定额发票 ----
        if re.search(r'定额发票|监制(?:章|机关)', text) and len(text) < 400:
            result['invoice_type'] = '定额发票'
            m = re.search(r'[¥￥]?\s*(\d+(?:\.\d{1,2})?)\s*元', text)
            if m:
                result['amount'] = f"{float(m.group(1)):.2f}"
            return True

        return False

    def _extract_fields(self, text: str) -> dict:
        """字段提取 - 改进版，优先用标签，备选用通用模式"""
        # 第一步：修复 OCR 拆行 & 统一标签
        text = self._normalize_text(text)

        result = {
            "date": None,
            "invoice_number": None,
            "buyer": None,
            "supplier": None,
            "amount": None,        # 价税合计
            "tax_free_amount": None,  # 合计金额（不含税）
            "tax_amount": None,    # 合计税额
            "invoice_type": None,  # 发票类型（特殊发票用）
        }

        # 第二步：特殊发票类型优先处理
        special_done = self._extract_special_invoice(text, result)

        # === 日期 ===
        # 优先匹配有明确标签的日期，再匹配年月日格式，最后才用紧凑数字（严格校验避免误匹配发票号）
        date_patterns = [
            # 有标签：开票日期 / 日期 后面跟年月日
            (r'(?:开票日期|日期)[：:\s]*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})', 1, 2, 3),
            # 标准年月日
            (r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})', 1, 2, 3),
            # 分隔符格式 2026-06-04 / 2026/06/04
            (r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})', 1, 2, 3),
            # 紧凑格式：严格要求月份01-12，日期01-31，且前后不是数字（避免匹配发票号）
            (r'(?<!\d)(20\d{2})(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])(?!\d)', 1, 2, 3),
        ]
        for entry in date_patterns:
            pattern, gi, gm, gd = entry
            match = re.search(pattern, text)
            if match:
                try:
                    y = int(match.group(gi))
                    m = int(match.group(gm))
                    d = int(match.group(gd))
                    if 1900 <= y <= 2100 and 1 <= m <= 12 and 1 <= d <= 31:
                        result["date"] = f"{y:04d}-{m:02d}-{d:02d}"
                        break
                except:
                    pass

        # === 发票号 ===
        # 支持多种格式：纯数字、含字母（如 O 代替 0）、含括号等
        inv_patterns = [
            r'(?:发票号|号码)[：\s]*([A-Z0-9\)\(]{15,})',  # 发票号标签 + 字母数字
            r'(?:发票号|号码)[：\s]*([0-9\)\(]{10,})',  # 发票号标签 + 纯数字（可能含括号）
            r'[A-Z0-9]{15,}',  # 任意15+位字母数字
            r'\d{15,25}',  # 任意15-25位纯数字（支持更长的全电发票号）
        ]
        for pattern in inv_patterns:
            if 'A-Z' in pattern or '(' in pattern or ')' in pattern:
                match = re.search(pattern, text)
                if match:
                    raw = match.group(1) if '(' in pattern else match.group(0)
                    clean = re.sub(r'[^\dA-Z0-9]', '', raw)
                    if 13 <= len(clean) <= 25:
                        result["invoice_number"] = clean
                        break
            else:
                matches = re.findall(pattern, text)
                if matches:
                    result["invoice_number"] = matches[0]
                    break

        # === 购买方和销售方 ===
        # 只是后缀、不含实质性公司名前缀 → 视为提取失败
        _SUFFIX_ONLY = re.compile(
            r'^(?:有限(?:责任)?公司|股份有限公司|集团公司|有限公司|责任公司|公司)$'
        )
        # 发票字段标签词本身不能作为公司名
        _LABEL_WORDS = frozenset({
            '名称', '金额', '税额', '地址', '电话', '合计', '税率',
            '备注', '开票人', '识别号', '统一社会', '纳税人', '规格',
            '项目', '单位', '数量', '单价', '备注', '信息',
        })

        # 政府机关/税务机关关键词 → 不应作为购买方/销售方
        _GOVT_WORDS = re.compile(
            r'税务[局所]|国家税务|地方税务|稽查局|国税局|地税局|财政局|监察局|'
            r'市场监督|行政管理局|公安局|政府|监制机关|主管税务'
        )

        def clean_company(name):
            """清理公司名：截断噪音、校验合法性"""
            if not name:
                return None
            name = name.strip()
            # 同行存在另一方标签时截断（如 "A公司 销售方名称: B公司"）
            name = re.split(r'(?:销售方|购买方)\s*名称', name)[0]
            # 截断在纳税人/地址/开户/电话/监制机关等发票字段标签处
            name = re.split(r'(?:纳税人|识别号|地址[、，,]|开户|电话|统一社会|监制机关|主管税务)', name)[0]
            # 去掉末尾长数字串（税号等）
            name = re.sub(r'\s*\d{8,}.*$', '', name)
            # 去掉括号内的企业类型说明（如"(个体工商户)"、"(个人)"）
            name = re.sub(r'\s*[（(]\s*(?:个体工商户|个人独资|自然人|个人)\s*[）)]', '', name)
            # 去掉末尾标点/空白
            name = re.sub(r'[：:\s，,。.]+$', '', name).strip()
            # 必须含中文
            if not re.search(r'[\u4e00-\u9fa5]', name):
                return None
            if len(name) < 2 or len(name) > 60:
                return None
            # 拒绝纯后缀（如 "有限公司"）
            if _SUFFIX_ONLY.match(name):
                return None
            # 拒绝发票字段标签词（如 "名称"）
            if name in _LABEL_WORDS:
                return None
            # 拒绝政府机关/税务机关（不应出现在购买方/销售方位置）
            if _GOVT_WORDS.search(name):
                return None
            return name

        def _multiline_company(text, label_pat):
            """
            在 text 中匹配 label_pat 后紧跟的公司名。
            用 [ \\t]* 不跨行；首行不完整时尝试拼接下一行。
            """
            m = re.search(label_pat + r'[ \t]*([^\n]*)', text, re.MULTILINE)
            if not m:
                return None
            first = m.group(1).strip()
            # 取下一行
            rest = text[m.end():]
            next_line = re.match(r'[ \t]*([^\n]{1,40})', rest)
            cont = next_line.group(1).strip() if next_line else ''

            # 判断是否需要拼接：首行缺省/很短/以不完整词结尾
            is_incomplete = (
                not first                                                        # 行内无内容
                or len(first) < 6                                                # 太短
                or re.search(r'(?:有限|股份|集团|科技|责任|管理|实业|发展)\s*$', first)
            )
            # 下一行像是公司后续（含公司关键词，且不是字段标签）
            is_continuation = bool(cont) and bool(
                re.search(r'(?:公司|有限|责任|集团|股份|管理|科技|发展|实业)', cont)
            ) and not re.search(
                r'(?:纳税人|识别号|地址|开户|电话|购买方|销售方|统一社会|信用代码)', cont
            )

            if is_incomplete and is_continuation:
                first = (first + cont).strip()
            elif not first and cont and not re.search(
                r'(?:纳税人|识别号|地址|开户|电话|购买方|销售方)', cont
            ):
                # 标签后为空行，直接取下一行内容
                first = cont

            return first or None

        # 策略0: 同行双名称（OCR 将两列合并到一行，如 "名称: A公司  名称: B公司"）
        same_line_both = re.search(
            r'[1l名]称[：:]\s*(.+?)\s{2,}[1l名]称[：:]\s*([^\n]+)', text)

        # 策略0b: "购买方名称: A公司 销售方名称: B公司" 同行完整格式
        explicit_both = re.search(
            r'购买方\s*名称[：:]\s*(.+?)\s+销售方\s*名称[：:]\s*([^\n]+)', text)

        # 策略1: 购买方——多模式，支持跨行，含OCR误读和替代标签
        buyer_raw = (
            _multiline_company(text, r'购买方\s*名称[：:]')
            or _multiline_company(text, r'(?:购买方|购\s*方|买\s*方)[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            # OCR 误读变体
            or _multiline_company(text, r'(?:构|沟)买方?[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'(?:构|沟)\s*方[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            # 替代标签：客户名称 / 付款方
            or _multiline_company(text, r'客户名称[：:]')
            or _multiline_company(text, r'付款方[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'付款方[：:]')
        )

        # 策略1: 销售方——多模式，支持跨行，含OCR误读和替代标签
        supplier_raw = (
            _multiline_company(text, r'销售方\s*名称[：:]')
            or _multiline_company(text, r'(?:销售方|销\s*方|卖\s*方)[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'销[^\n]{0,20}?[名1l]\s*[称秄][：:]')
            # OCR 误读变体：消售/晓售/销害/销室/销富
            or _multiline_company(text, r'(?:消|晓)售方?[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'(?:消|晓)\s*方[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'销[害室富][万放仿]?[^\n]{0,30}?[名1l]\s*[称][：:]')
            # 替代标签：收款方 / 供应方 / 卖方 / 经营方 / 服务提供方
            or _multiline_company(text, r'收款方[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'收款方[：:]')
            or _multiline_company(text, r'供应方[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'供应方[：:]')
            or _multiline_company(text, r'卖方[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'卖方[：:]')
            or _multiline_company(text, r'(?:经营方|服务提供方)[^\n]{0,30}?[名1l]\s*[称秄][：:]')
            or _multiline_company(text, r'(?:经营方|服务提供方)[：:]')
        )

        # 优先级：explicit_both > 策略1个别匹配 > same_line_both
        if explicit_both:
            if not buyer_raw:
                buyer_raw = explicit_both.group(1)
            if not supplier_raw:
                supplier_raw = explicit_both.group(2)

        # 仅在特殊发票未提前设置值时才写入（防止覆盖航空行程单等已设定的旅客/路线）
        if buyer_raw and not result["buyer"]:
            result["buyer"] = clean_company(buyer_raw)
        if supplier_raw and not result["supplier"]:
            result["supplier"] = clean_company(supplier_raw)

        # 策略0回填（同行双名称）
        if same_line_both and (not result["buyer"] or not result["supplier"]):
            b = clean_company(same_line_both.group(1))
            s = clean_company(same_line_both.group(2))
            if b and not result["buyer"]:
                result["buyer"] = b
            if s and not result["supplier"]:
                result["supplier"] = s

        # 策略2: 按出现顺序收集所有 "名称:" 内容，补全缺失的一方
        # 注意：用 [ \t]* 而非 \s* 避免跨行消耗（\s 会匹配换行符）
        if not result["buyer"] or not result["supplier"]:
            company_lines = []
            # 支持 "名称:" / "1称:" / "l称:" / "1秄:" 等OCR变体
            for pat in (r'[名1l]\s*[称秄][：:][ \t]*([^\n]+)', r'名称[：:][ \t]*([^\n]+)',
                        r'客户名称[：:][ \t]*([^\n]+)', r'收款方[：:][ \t]*([^\n]+)'):
                for m in re.findall(pat, text):
                    c = clean_company(m)
                    if c and c not in company_lines:
                        company_lines.append(c)
                if company_lines:
                    break

            # 策略3: 通用企业词尾识别
            # 用 finditer 带上下文，排除开户银行行/税务机关行
            if len(company_lines) < 2:
                # 先把开户行/开户银行所在行从文本里剔除，防止误匹配
                text_no_bank = re.sub(r'(?m)^[^\n]*(?:开户行|开户银行|银行账号|账号)[^\n]*$', '', text)
                # 也剔除纳税人识别号行和主管税务行
                text_no_bank = re.sub(r'(?m)^[^\n]*(?:纳税人识别号|统一社会信用代码|主管税务)[^\n]*$', '', text_no_bank)
                for mat in re.finditer(
                    r'[\u4e00-\u9fa5]{2,}(?:公司|有限|分公司|集团|股份|企业|'
                    r'研究所|医院|学校|协会|中心|院|所|厂|部)', text_no_bank
                ):
                    c = clean_company(mat.group(0))
                    if c and c not in company_lines:
                        company_lines.append(c)

            # 策略3b: 如果仍然不够2个公司，尝试从整段文本中扫描
            # 住宿发票中销售方经常是酒店名（如"XX酒店"、"XX宾馆"、"XX饭店"）
            if len(company_lines) < 2:
                for mat in re.finditer(
                    r'[\u4e00-\u9fa5]{2,}(?:酒店|宾馆|饭店|旅馆|招待所|公寓|'
                    r'出租(?:汽车)?公司|客运公司)', text_no_bank
                ):
                    c = clean_company(mat.group(0))
                    if c and c not in company_lines:
                        company_lines.append(c)

            if not result["buyer"] and len(company_lines) >= 1:
                result["buyer"] = company_lines[0]
            if not result["supplier"]:
                buyer_val = result.get("buyer")
                for c in company_lines:
                    if c != buyer_val:
                        result["supplier"] = c
                        break

        def _find_amount(patterns, text):
            """通用金额提取，返回匹配到的第一个有效金额字符串"""
            for pattern in patterns:
                matches = re.findall(pattern, text)
                if matches:
                    try:
                        vals = []
                        for m in matches:
                            if isinstance(m, tuple):
                                vals.append(f"{m[0]}.{m[1]}")
                            else:
                                vals.append(str(m))
                        best = str(max(float(v) for v in vals if v))
                        # 格式化为两位小数
                        return f"{float(best):.2f}"
                    except:
                        pass
            return None

        # === 价税合计（总金额）===
        # 仅在特殊发票未提前设置金额时才运行通用提取（防止覆盖已识别值）
        if not result["amount"]:
            total_patterns = [
                # 标准VAT发票：价税合计(小写) ¥XXX.XX
                r'价税合计[^0-9\n]{0,10}小写[）)]*\s*[垒¥￥垩圓Y]?\s*([0-9]{1,10}\.[0-9]{2})',
                r'小写[）)]*\s*[垒¥￥垩圓Y]?\s*([0-9]{1,10}\.[0-9]{2})',
                r'价税合计[^0-9\n]{0,20}([0-9]{1,10}\.[0-9]{2})',
                # 航空/出租/住宿：合计/金额/票价标签
                r'(?:合计|实付|应付|票价|金额)[：:\s]*[¥￥]?\s*([0-9]{1,10}\.[0-9]{2})',
                r'(?:合计|实付|应付|票价|金额)[：:\s]*[¥￥]?\s*([0-9]{1,6}(?:\.[0-9]{1,2})?)',
                # ¥符号后跟金额（含一位小数）
                r'[¥￥垩圓Y垒]\s*([0-9]{1,10}(?:\.[0-9]{1,2})?)',
                # 定额发票：直接写 "XX元" 或 "¥XX"（无小数点）
                r'(?<![0-9])([0-9]{1,6})\s*元(?:整)?(?![0-9])',
                # 最后兜底：最大的两位小数
                r'([0-9]{1,10}\.[0-9]{2})',
            ]
            result["amount"] = _find_amount(total_patterns, text)

        # === 合计金额（不含税）+ 税额 ===
        # OCR 经常丢失¥符号或将其误读为"半/垩/垒"，也经常拆散表格行。
        # 最可靠的方法：找所有小数，找两个加起来约等于价税合计的组合。

        # 策略1：标签直接匹配（有标签时最精准）
        amtpat = r'[¥￥垒垩圓Y半]?\s*([0-9]{1,10}\.[0-9]{2})'
        # 合计行两列同行
        two_num = re.search(r'合计\s*' + amtpat + r'[\s\n]+' + amtpat, text)
        if two_num:
            result["tax_free_amount"] = f"{float(two_num.group(1)):.2f}"
            result["tax_amount"]      = f"{float(two_num.group(2)):.2f}"
        else:
            # 带税额标签
            m = re.search(r'(?:合计税额|税\s*额)[：:\s]*' + amtpat, text)
            if m:
                result["tax_amount"] = f"{float(m.group(1)):.2f}"
            m = re.search(r'(?:不含税|合计金额)[：:\s]*' + amtpat, text)
            if m:
                result["tax_free_amount"] = f"{float(m.group(1)):.2f}"

        # 策略2：若标签匹配失败，用"配对加法"推断
        # 找所有小数（排除价税合计本身），看哪两个之和约等于总金额
        if result["amount"] and (not result["tax_free_amount"] or not result["tax_amount"]):
            try:
                total = float(result["amount"])
                # 提取文本中所有 x.xx 格式小数，去掉和总价相同的
                candidates = list(dict.fromkeys(
                    float(m) for m in re.findall(r'\b(\d{1,8}\.\d{2})\b', text)
                    if abs(float(m) - total) > 0.01 and float(m) > 0
                ))
                found = False
                for i, a in enumerate(candidates):
                    for b in candidates[i+1:]:
                        if abs(a + b - total) <= 0.05:   # 允许0.05元的舍入误差
                            bigger  = max(a, b)
                            smaller = min(a, b)
                            result["tax_free_amount"] = f"{bigger:.2f}"
                            result["tax_amount"]      = f"{smaller:.2f}"
                            found = True
                            break
                    if found:
                        break
            except:
                pass

        # 策略3：只知道其中一个，推算另一个
        if result["amount"] and result["tax_free_amount"] and not result["tax_amount"]:
            try:
                tax = round(float(result["amount"]) - float(result["tax_free_amount"]), 2)
                if tax > 0:
                    result["tax_amount"] = f"{tax:.2f}"
            except:
                pass
        if result["amount"] and result["tax_amount"] and not result["tax_free_amount"]:
            try:
                base = round(float(result["amount"]) - float(result["tax_amount"]), 2)
                if base > 0:
                    result["tax_free_amount"] = f"{base:.2f}"
            except:
                pass

        return result


def generate_filename(data: dict, original_ext: str) -> str:
    """生成发票新文件名：日期_销售方_购买方_金额元.ext"""
    date     = data.get('date') or '0000-01-01'
    buyer    = (data.get('buyer')    or '')[:20]
    supplier = (data.get('supplier') or '')[:20]
    amount   = data.get('amount') or '0.00'

    new_name = f"{date}_{supplier}_{buyer}_{amount}元{original_ext}"
    new_name = re.sub(r'[\\/:*?"<>|]', '', new_name)
    new_name = re.sub(r'_+', '_', new_name).strip('_')
    return new_name or f"invoice_{datetime.now().strftime('%Y%m%d%H%M%S')}{original_ext}"


def generate_train_filename(data: dict, original_ext: str) -> str:
    """生成火车票新文件名: 日期_出发站-到达站_票价元.ext"""
    date    = data.get('date') or '0000-01-01'
    from_st = (data.get('from_station') or '')[:10]
    to_st   = (data.get('to_station')   or '')[:10]
    price   = data.get('price') or '0.00'
    route   = f"{from_st}-{to_st}" if (from_st or to_st) else ''
    parts   = [p for p in [date, route, f"{price}元"] if p]
    new_name = '_'.join(parts) + original_ext
    new_name = re.sub(r'[\\/:*?"<>|]', '', new_name)
    new_name = re.sub(r'_+', '_', new_name).strip('_')
    return new_name or f"train_{datetime.now().strftime('%Y%m%d%H%M%S')}{original_ext}"


# ======================== 火车票提取 ========================

# T3出行/网约车关键词集合（用于类型检测，优先级高于火车票）
_T3_KEYWORDS = re.compile(
    r'T3\s*出行|T3出行|滴滴出行|滴滴打车|曹操出行|高德打车|美团打车|'
    r'花小猪|首汽约车|享道出行|如祺出行|阳光出行|万顺叫车|'
    r'网约车|电子行程单|出行服务|打车电子发票|出租汽车电子发票|'
    r'运客(?:费|服务)|客运(?:服务|费)|运输服务|出行平台'
)


def _abbreviate_city(name: str) -> str:
    """城市名简写：取前2个汉字（如"上海虹桥"→"上海"，"北京南站"→"北京"）"""
    if not name:
        return ''
    # 去掉常见后缀
    name = re.sub(r'(?:市|区|县|站|机场|虹桥|南站|北站|东站|西站)$', '', name)
    # 取前2个汉字
    m = re.match(r'([\u4e00-\u9fa5]{1,2})', name)
    return m.group(1) if m else name[:2]


def generate_t3_filename(data: dict, original_ext: str) -> str:
    """生成T3出行/网约车发票新文件名: 日期_出发地简写-到达地简写_金额元.ext"""
    date    = data.get('date') or '0000-01-01'
    from_city = _abbreviate_city(data.get('from_station') or data.get('from_city') or '')
    to_city   = _abbreviate_city(data.get('to_station') or data.get('to_city') or '')
    price   = data.get('price') or data.get('amount') or '0.00'

    route = f"{from_city}-{to_city}" if (from_city or to_city) else ''
    parts = [p for p in [date, route, f"{float(price):.2f}元"] if p]
    new_name = '_'.join(parts) + original_ext
    new_name = re.sub(r'[\\/:*?"<>|]', '', new_name)
    new_name = re.sub(r'_+', '_', new_name).strip('_')
    return new_name or f"t3_{datetime.now().strftime('%Y%m%d%H%M%S')}{original_ext}"


# 火车票关键词集合（用于类型检测）
_TRAIN_KEYWORDS = re.compile(
    r'车\s*次|检\s*票|候\s*车|动\s*车|高\s*铁|火\s*车\s*票|硬\s*卧|软\s*卧|硬\s*座|'
    r'二\s*等\s*座|一\s*等\s*座|商\s*务\s*座|无\s*座|出\s*发\s*站|到\s*达\s*站|'
    r'网络购票|铁路电子客票|中国铁路|12306|票价[：:\s]*[¥￥]?\d|'
    r'列\s*车\s*号|乘\s*车\s*日|席\s*别|始\s*发\s*站|终\s*到\s*站|补\s*票|'
    r'开\s*车\s*时\s*间|出\s*发\s*时\s*间|铁\s*路\s*客\s*票'
)
# 修复：去掉 \b，在中文环境下 \b 不能正确匹配（中文字符属于 \w，无法形成边界）
_TRAIN_NUMBER_RE = re.compile(r'(?<![A-Z\d])([GDTZKCY]\d{1,4})(?!\d)')

# 合同关键词（用于类型检测）
_CONTRACT_PARTY_A  = re.compile(r'甲\s*方|买\s*方|委\s*托\s*方|发\s*包\s*方|采\s*购\s*方|招\s*标\s*人')
_CONTRACT_PARTY_B  = re.compile(r'乙\s*方|卖\s*方|承\s*包\s*方|承\s*接\s*方|供\s*货\s*方|中\s*标\s*人')
_CONTRACT_STRONG   = re.compile(
    r'本\s*合\s*同|本\s*协\s*议|合\s*同\s*编\s*号|甲\s*乙\s*双\s*方|买\s*卖\s*双\s*方|'
    r'委\s*托\s*方|发\s*包\s*方|合\s*同\s*金\s*额|合\s*同\s*总\s*额|合\s*同\s*总\s*价|'
    r'平\s*等\s*自\s*愿|协\s*商\s*一\s*致|合\s*同\s*协\s*议\s*书|货\s*物\s*采\s*购\s*合\s*同|'
    r'采\s*购\s*合\s*同|服\s*务\s*合\s*同|工\s*程\s*合\s*同|建\s*设\s*工\s*程\s*合\s*同'
)


def detect_doc_type(text: str) -> str:
    """根据 OCR 文本判断是火车票/合同/发票，返回 't3'/'train'/'contract'/'invoice'"""
    # 0. T3出行/网约车（优先级最高，避免被误判为火车票）
    if _T3_KEYWORDS.search(text):
        return 't3'
    # 1. 强火车票关键词（优先级最高）
    if _TRAIN_KEYWORDS.search(text):
        return 'train'
    # 2. 合同强关键词（优先于弱车次号检测，防止项目编号被误判为车次）
    has_a = bool(_CONTRACT_PARTY_A.search(text))
    has_b = bool(_CONTRACT_PARTY_B.search(text))
    has_s = bool(_CONTRACT_STRONG.search(text))
    if has_s or (has_a and has_b):
        return 'contract'
    # 3. 弱车次号检测（仅在无合同信号时才判断）
    if _TRAIN_NUMBER_RE.search(text):
        return 'train'
    return 'invoice'


class TrainTicketExtractor:
    """火车票 OCR 字段提取"""

    # 站名结尾词（帮助识别站名）
    _STATION_SUFFIX = re.compile(r'[\u4e00-\u9fa5]{2,8}(?:站|虹桥|南|北|东|西|高铁)?')
    # 座位类型
    _SEAT_TYPES = ['商务座', '特等座', '一等座', '二等座', '软卧上', '软卧下', '硬卧上', '硬卧中',
                   '硬卧下', '软卧', '硬卧', '硬座', '无座', '动卧']

    def __init__(self, reader):
        self.reader = reader

    # 复用 InvoiceExtractor 的图像读取方法（避免重复）
    def _pdf_to_image(self, pdf_path):
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=200)
            if not images:
                raise Exception("PDF 无法转换")
            return cv2.cvtColor(np.array(images[0]), cv2.COLOR_RGB2BGR)
        except ImportError:
            raise Exception("缺少 pdf2image 库")
        except Exception as e:
            raise Exception(f"PDF 转图片失败: {e}")

    def _image_file_to_array(self, image_path):
        try:
            image = Image.open(image_path).convert('RGB')
            return cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        except Exception as e:
            raise Exception(f"图片读取失败: {e}")

    def _preprocess(self, image):
        try:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            return cv2.cvtColor(clahe.apply(gray), cv2.COLOR_GRAY2BGR)
        except:
            return image

    def extract(self, file_path: str) -> dict:
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        if ext == '.pdf':
            img = self._pdf_to_image(str(file_path))
        else:
            img = self._image_file_to_array(str(file_path))
        if img is None or img.size == 0:
            raise Exception("无法读取图像")
        img = self._preprocess(img)
        results = self.reader.readtext(img, detail=0)
        if not results:
            raise Exception("OCR 无法识别")
        text = '\n'.join(results)
        fields = self._extract_fields(text, str(file_path.stem))
        fields['_raw_text'] = text
        return fields

    # 不应被识别为站名的元词
    _STATION_BLACKLIST = re.compile(
        r'^(?:出发站|到达站|始发站|终到站|目的地|经由|中转|检票口|候车|开车)$')

    def _clean_station(self, name: str) -> str:
        """清理站名：去掉"站"后缀、过滤黑名单词"""
        name = name.strip()
        # 过滤黑名单元词
        if self._STATION_BLACKLIST.match(name):
            return ''
        # 去掉末尾的"站"（保留内嵌的，如"北站"本身不去）
        if name.endswith('站') and len(name) > 2:
            name = name[:-1]
        return name

    def _extract_fields(self, text: str, stem: str = '') -> dict:
        result = {
            'date':           None,
            'train_number':   None,
            'from_station':   None,
            'to_station':     None,
            'passenger_name': None,
            'seat':           None,
            'seat_type':      None,
            'price':          None,
            'depart_time':    None,
        }

        # === 车次 ===
        # 优先：有标签（车次:/列车号:）
        tn_labeled = re.search(
            r'(?:车\s*次|列\s*车\s*号)[：:\s]*([GDTZKCY]\d{1,4})', text)
        if tn_labeled:
            result['train_number'] = tn_labeled.group(1)
        else:
            # 修复了的正则（去掉 \b，用显式边界）
            m = _TRAIN_NUMBER_RE.search(text)
            if m:
                result['train_number'] = m.group(1)

        # === 日期 ===
        # 注意：铁路电子客票有"乘车日期"（乘车日）和"开票日期"（开票日），取乘车日期
        # 先把"开票日期"这行从文本中临时遮掉，避免误匹配
        text_no_invoice_date = re.sub(r'开\s*票\s*日\s*期[：:\s]*\d{4}年\d{1,2}月\d{1,2}日', '', text)
        date_pats = [
            # 优先：有明确标签的乘车/出发日期
            (r'(?:乘车日期|出发日期|乘\s*车\s*日)[：:\s]*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})', 2, 3, 4),
            # 次选：任意"XXXX年XX月XX日"（已排除开票日期行）
            (r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})', 1, 2, 3),
            # ISO 格式
            (r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})', 1, 2, 3),
        ]
        for pat, gi, gm, gd in date_pats:
            src = text if gi == 2 else text_no_invoice_date  # 有标签时用全文，否则用遮掉版本
            m = re.search(pat, src)
            if m:
                try:
                    y, mo, d = int(m.group(gi)), int(m.group(gm)), int(m.group(gd))
                    if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
                        result['date'] = f'{y:04d}-{mo:02d}-{d:02d}'
                        break
                except Exception:
                    pass

        # === 出发时间 ===
        tm = re.search(r'(?:出发时间|开车时间|发车)[：:\s]*(\d{1,2}:\d{2})', text)
        if tm:
            result['depart_time'] = tm.group(1)
        else:
            tm = re.search(r'(\d{2}):(\d{2})(?::(\d{2}))?', text)
            if tm:
                result['depart_time'] = tm.group(0)[:5]

        # === 乘客姓名（先提取，用于后续站名过滤）===
        _NAME_BLACKLIST = {
            '出发', '到达', '乘坐', '车次', '购票', '旅客', '列车', '中国',
            '铁路', '上海', '北京', '广州', '深圳', '成都', '武汉', '南京',
            '高铁', '动车', '候车', '检票', '开车', '席别', '座位', '票价',
        }
        name_pats = [
            # 有标签
            r'(?:姓\s*名|旅\s*客|购\s*票\s*人|乘\s*客)[：:\s]*([\u4e00-\u9fa5]{2,4})',
            # 身份证/护照关键字前
            r'([\u4e00-\u9fa5]{2,4})[（\(]?(?:居民身份证|身份证|护照)',
            # 身份证号码前（全数字 15-18 位）
            r'([\u4e00-\u9fa5]{2,4})\s*\d{15,18}[Xx]?',
            # 隐码身份证前（含 * 星号遮盖）
            r'([\u4e00-\u9fa5]{2,4})\s*\*{4,}',
            # 铁路电子客票：隐码ID（****5127）后换行即为姓名
            r'\*{4,}\d+\n([\u4e00-\u9fa5]{2,4})',
            # 隐码ID或长数字串后换行（更宽松兜底）
            r'(?:\*{4,}|\d{6,})[^\n]*\n([\u4e00-\u9fa5]{2,4})\n',
            # 姓名出现在座位类型之前（如 "张三 二等座"）
            r'([\u4e00-\u9fa5]{2,4})\s+(?:商务座|特等座|一等座|二等座|软卧|硬卧|硬座|无座)',
            # 姓名出现在票价/¥ 之前（如 "张三 ¥553.50"）
            r'([\u4e00-\u9fa5]{2,4})\s*[¥￥]\s*\d',
        ]
        for p in name_pats:
            nm = re.search(p, text)
            if nm:
                name = nm.group(1).strip()
                if name not in _NAME_BLACKLIST:
                    result['passenger_name'] = name
                    break

        # 站名过滤时动态排除已识别的乘客姓名，防止人名末字含方位词被误判为站名
        _station_extra_excl = set()
        if result['passenger_name']:
            _station_extra_excl.add(result['passenger_name'])

        def _clean_st(name: str) -> str:
            s = self._clean_station(name)
            return '' if s in _station_extra_excl else s

        # === 出发站 / 到达站 ===
        # 辅助：在标签后（同行或跨行，跳过拼音行）提取站名
        def _find_labeled_station(label_regex):
            """支持三种格式：
            同行:   始发站：上海虹桥
            跨行:   始发站\n上海虹桥
            含拼音: 始发站\nShanghaihongqiao\n上海虹桥
            """
            # 同行格式
            m = re.search(label_regex + r'[ \t]*([\u4e00-\u9fa5]{2,12})', text)
            if m:
                s = _clean_st(m.group(1))
                if s:
                    return s
            # 跨行：向后扫描最多 4 行，跳过纯英文（拼音）行
            m2 = re.search(label_regex, text)
            if m2:
                after = text[m2.end():]
                for line in after.split('\n')[:4]:
                    line = line.strip()
                    if not line:
                        continue
                    # 跳过纯拼音/英文行
                    if re.match(r'^[A-Za-z0-9\s\-]+$', line):
                        continue
                    ch = re.search(r'([\u4e00-\u9fa5]{2,12})', line)
                    if ch:
                        s = _clean_st(ch.group(1))
                        if s:
                            return s
            return None

        # 策略1a：明确标签（支持同行/跨行/含拼音跨行）
        _from_labeled = _find_labeled_station(r'(?:出\s*发\s*站|始\s*发\s*站)')
        _to_labeled   = _find_labeled_station(r'(?:到\s*达\s*站|终\s*到\s*站|目\s*的\s*地)')
        if _from_labeled: result['from_station'] = _from_labeled
        if _to_labeled:   result['to_station']   = _to_labeled

        # 策略1b：双栏格式「始发站  终到站」→ 向后扫描，跳过拼音行，取第一个含两组中文的行
        if not result['from_station'] or not result['to_station']:
            two_col_hdr = re.search(r'(?:始发站|出发站)[ \t]+(?:终到站|到达站)', text)
            if two_col_hdr:
                after = text[two_col_hdr.end():]
                for line in after.split('\n')[:5]:
                    line = line.strip()
                    if not line:
                        continue
                    if re.match(r'^[A-Za-z0-9\s\-]+$', line):
                        continue  # 跳过拼音行
                    parts = re.findall(r'[\u4e00-\u9fa5]{2,12}', line)
                    if len(parts) >= 2:
                        f = _clean_st(parts[0])
                        t = _clean_st(parts[1])
                        if f and not result['from_station']: result['from_station'] = f
                        if t and not result['to_station']:   result['to_station']   = t
                        break

        # 策略1c：单标签行后紧跟两组中文（紧凑双栏）
        if not result['from_station'] or not result['to_station']:
            compact = re.search(
                r'(?:始发站|出发站|出发)[^\n]*\n'
                r'[ \t]*([\u4e00-\u9fa5]{2,10})\s+([\u4e00-\u9fa5]{2,10})',
                text)
            if compact:
                f = _clean_st(compact.group(1))
                t = _clean_st(compact.group(2))
                if f and not result['from_station']: result['from_station'] = f
                if t and not result['to_station']:   result['to_station']   = t

        # 策略1d：铁路电子客票标准格式 —— 出发站\n到达站\nG/D/K/Z/T/C-XXXX
        # 在车次号之前的行里，反向扫描找到紧邻的中文站名行
        if not result['from_station'] or not result['to_station']:
            _tn = result.get('train_number') or ''
            if _tn:
                _m_t = re.search(r'(?<![A-Z\d])' + re.escape(_tn) + r'(?!\d)', text)
                if _m_t:
                    _before = text[:_m_t.start()]
                    _lines = [l.strip() for l in _before.split('\n') if l.strip()]
                    _stn_lines: list = []
                    for _ln in reversed(_lines[-10:]):
                        # 停止条件：遇到发票元数据行
                        if re.search(r'(?:发票号码|开票日期|全国|监制|税务总局|统一发票)', _ln):
                            break
                        # 纯英文/数字行跳过
                        if re.match(r'^[A-Za-z0-9\s\'./:\-]+$', _ln):
                            continue
                        # 匹配纯中文站名（2-8字，可选末尾"站"）
                        if re.match(r'^[\u4e00-\u9fa5]{2,8}站?$', _ln):
                            _stn_lines.insert(0, _ln)  # 还原文本顺序
                            if len(_stn_lines) >= 2:
                                break
                    if len(_stn_lines) >= 2:
                        f = _clean_st(_stn_lines[0])
                        t = _clean_st(_stn_lines[1])
                        if f and not result['from_station']: result['from_station'] = f
                        if t and not result['to_station']:   result['to_station']   = t
                    elif len(_stn_lines) == 1:
                        s = _clean_st(_stn_lines[0])
                        if s and not result['from_station']: result['from_station'] = s

        # 策略2：箭头/横线分隔（上海虹桥→北京南）
        if not result['from_station'] or not result['to_station']:
            arrow = re.search(
                r'([\u4e00-\u9fa5]{2,10}(?:站)?)\s*[→➜>—至]\s*([\u4e00-\u9fa5]{2,10}(?:站)?)',
                text)
            if arrow:
                f = _clean_st(arrow.group(1))
                t = _clean_st(arrow.group(2))
                if f and not result['from_station']: result['from_station'] = f
                if t and not result['to_station']:   result['to_station']   = t

        # 策略3：含复合方位词后缀（东站/西站/南站/北站/虹桥/高铁）
        if not result['from_station'] or not result['to_station']:
            station_candidates = re.findall(
                r'([\u4e00-\u9fa5]{2,8}(?:虹桥|高铁|北站|南站|东站|西站)(?:站)?'
                r'|[\u4e00-\u9fa5]{4,8}(?:南|北|东|西)(?:站)?)',
                text)
            clean_cands = []
            for sc in station_candidates:
                s = _clean_st(sc)
                if s and s not in clean_cands:
                    clean_cands.append(s)
            if not result['from_station'] and len(clean_cands) >= 1:
                result['from_station'] = clean_cands[0]
            if not result['to_station'] and len(clean_cands) >= 2:
                result['to_station'] = clean_cands[1]

        # 策略4：带"站"字后缀的任意站名（兜底）—— 去重后按序取两个不同站
        if not result['from_station'] or not result['to_station']:
            _seen4: set = set()
            with_zhan = []
            for s in [_clean_st(x) for x in re.findall(r'([\u4e00-\u9fa5]{2,8}站)', text)]:
                if s and s not in _seen4:
                    _seen4.add(s)
                    with_zhan.append(s)
            _already = result['from_station']
            if not result['from_station'] and with_zhan:
                result['from_station'] = with_zhan[0]
            if not result['to_station']:
                for s in with_zhan:
                    if s != result['from_station']:
                        result['to_station'] = s
                        break

        # 策略5：纯汉字两段式，最后兜底（仅当两站都缺失且含地名特征词）
        if not result['from_station'] and not result['to_station']:
            _excl_names = _station_extra_excl | {result.get('passenger_name') or ''}
            _seen5: set = set()
            pure_zh = []
            for s in re.findall(r'([\u4e00-\u9fa5]{2,6})', text):
                if (s not in _excl_names
                        and not self._STATION_BLACKLIST.match(s)
                        and re.search(r'(?:站|桥|路|门|场|港)', s)
                        and s not in _seen5):
                    _seen5.add(s)
                    pure_zh.append(_clean_st(s))
            pure_zh = [s for s in pure_zh if s]
            if pure_zh:
                result['from_station'] = pure_zh[0]
            if len(pure_zh) >= 2:
                result['to_station'] = pure_zh[1]

        # 最终校验：出发==到达 → 说明兜底策略出错，清空出发站
        if (result['from_station'] and result['to_station']
                and result['from_station'] == result['to_station']):
            result['from_station'] = None

        # === 座位类型 ===
        for st in self._SEAT_TYPES:
            if st in text:
                result['seat_type'] = st
                break
        # 电子客票"席别"标签
        if not result['seat_type']:
            xi = re.search(r'席\s*别[：:\s]*([\u4e00-\u9fa5]{2,5})', text)
            if xi:
                result['seat_type'] = xi.group(1)

        # === 座位号 ===
        seat_pats = [
            r'(\d{1,2}\s*车\s*\d{1,2}\s*[A-F号])',   # 08车06A号
            r'([A-F]\d\s*车厢?\s*\d{1,2}\s*号?)',      # A3车厢6号
            r'(\d{1,2}[A-F]\d?)',                      # 紧凑格式 06A
        ]
        for sp in seat_pats:
            sm = re.search(sp, text)
            if sm:
                result['seat'] = sm.group(1).strip()
                break

        # === 票价 ===
        price_pats = [
            r'(?:票\s*价|价\s*格|金\s*额|票\s*款)[：:\s]*[¥￥]?\s*(\d+\.?\d*)',
            r'[¥￥]\s*(\d+\.?\d*)',
            r'(\d{2,5}\.\d{1,2})\s*元',
            r'合\s*计[：:\s]*[¥￥]?\s*(\d+\.?\d*)',
        ]
        for pp in price_pats:
            pm = re.search(pp, text)
            if pm:
                try:
                    result['price'] = f"{float(pm.group(1)):.2f}"
                    break
                except ValueError:
                    pass
        # 兜底：文中最大小数
        if not result['price']:
            decimals = [float(x) for x in re.findall(r'(?<!\d)(\d{1,5}\.\d{1,2})(?!\d)', text)
                        if 1 <= float(x) <= 10000]
            if decimals:
                result['price'] = f"{max(decimals):.2f}"

        # === 从文件名补充缺失字段 ===
        if stem:
            parts = re.split(r'[_\-\s]', stem)
            for part in parts:
                tn = re.match(r'^([GDTZKCY]\d{1,4})$', part)
                if tn and not result['train_number']:
                    result['train_number'] = tn.group(1)

        return result


# ======================== 合同提取 ========================

def _abbreviate_party(name: str) -> str:
    """从公司/单位名称提取 4-6 个关键字（去除括号地名和通用后缀）"""
    if not name:
        return ''
    # 去除括号内的地名（如"(苏州)"、"（上海）"）
    name = re.sub(r'[（(][^）)]{1,8}[）)]', '', name).strip()
    name = re.sub(r'\s+', '', name)

    # 按从长到短的顺序尝试剥离法律后缀（避免短后缀先匹配导致漏剥）
    _SUFFIXES = [
        '有限责任公司', '股份有限公司', '集团有限公司', '集团公司',
        '总公司', '分公司', '有限公司',
    ]
    for sfx in _SUFFIXES:
        if name.endswith(sfx):
            candidate = name[:-len(sfx)]
            if len(candidate) >= 2:   # 剥离后名称至少保留2字
                name = candidate
            break

    # 若仍超6字，取前6字
    if len(name) > 6:
        name = name[:6]
    return name


class ContractExtractor:
    """合同 OCR 字段提取（支持甲方/乙方 和 买方/卖方 等多种格式）"""

    def _extract_fields(self, text: str) -> dict:
        result = {
            'contract_name': '',
            'sign_date':     '',
            'party_a':       '',
            'party_b':       '',
            'amount':        '',
        }
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        self._extract_contract_name(lines, text, result)
        self._extract_sign_date(lines, text, result)
        self._extract_parties(lines, text, result)
        self._extract_amount(lines, text, result)
        return result

    # ---- 合同名称 ----
    def _extract_contract_name(self, lines, text, result):
        # 1. 标注字段
        for line in lines:
            m = re.search(r'(?:合同名称|协议名称|项目名称)[：:]\s*(.{2,40})', line)
            if m:
                name = re.sub(r'\s+', '', m.group(1)).strip()
                name = re.sub(r'[（(].*?[）)]', '', name).strip()
                if name and len(name) >= 2:
                    result['contract_name'] = name[:20]
                    return
        # 2. 标题行（前10行，含"合同"/"协议"词）
        for line in lines[:10]:
            clean = re.sub(r'[《》【】\[\]（(）)\s]+', '', line)
            if 2 < len(clean) <= 20 and re.search(r'合同|协议书', clean):
                if clean not in ('合同', '协议书', '协议', '本合同', '本协议'):
                    result['contract_name'] = clean
                    return
        # 3. 全文最先出现的 "XX合同" / "XX协议书"
        m = re.search(r'[\u4e00-\u9fff]{2,12}(?:采购合同|服务合同|工程合同|合同|协议书)', text)
        if m:
            result['contract_name'] = m.group(0)[:20]

    # ---- 签订日期 ----
    def _extract_sign_date(self, lines, text, result):
        pats = [
            r'(?:签订|签署|签约|合同)\s*[日期时间]*\s*[：:\s]+(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
            r'于\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
            r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
            r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
        ]
        for pat in pats:
            m = re.search(pat, text)
            if m:
                result['sign_date'] = (
                    f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
                )
                return

    # ---- 甲方/乙方（买方/卖方） ----
    def _extract_parties(self, lines, text, result):
        # 模式：关键词 + 可选冒号/空格 + 名称
        A_KW = r'(?:甲\s*方|买\s*方|委\s*托\s*方|发\s*包\s*方|采\s*购\s*方|招\s*标\s*人)'
        B_KW = r'(?:乙\s*方|卖\s*方|承\s*包\s*方|承\s*接\s*方|供\s*货\s*方|中\s*标\s*人)'
        # 支持 "买方 单位全称 XXX" 和 "买方：XXX" 两种格式
        a_pats = [
            A_KW + r'(?:\s*单位全称|\s*名称|\s*（盖章）|\s*\(盖章\))?\s*[：:\s]+([^\n]{2,35})',
        ]
        b_pats = [
            B_KW + r'(?:\s*单位全称|\s*名称|\s*（盖章）|\s*\(盖章\))?\s*[：:\s]+([^\n]{2,35})',
        ]

        def _clean(raw):
            val = re.sub(r'[（(][^）)]{1,8}[）)]', '', raw).strip()
            val = re.sub(r'\s+', '', val)
            # 去除捕获到的关键词前缀（防止把下一行关键词算进来）
            val = re.split(r'甲方|乙方|买方|卖方|承包方|委托方|招标人|中标人', val)[0]
            return val[:30]

        for line in lines:
            if not result['party_a']:
                for pat in a_pats:
                    m = re.search(pat, line)
                    if m:
                        v = _clean(m.group(1))
                        if len(v) >= 2:
                            result['party_a'] = v
                            break
            if not result['party_b']:
                for pat in b_pats:
                    m = re.search(pat, line)
                    if m:
                        v = _clean(m.group(1))
                        if len(v) >= 2:
                            result['party_b'] = v
                            break
            if result['party_a'] and result['party_b']:
                break

        # 兜底：宽松扫描（针对OCR排版混乱的情况）
        if not result['party_a']:
            m = re.search(A_KW + r'[\s：:]+([^\n（(]{2,30})', text)
            if m:
                result['party_a'] = _clean(m.group(1))
        if not result['party_b']:
            m = re.search(B_KW + r'[\s：:]+([^\n（(]{2,30})', text)
            if m:
                result['party_b'] = _clean(m.group(1))

    # ---- 合同金额 ----
    def _extract_amount(self, lines, text, result):
        amount_pats = [
            # 合同总额/合同金额/合同价 + 数字
            (r'(?:合同[总]?[额价款金]|总金额|总价款|合同总价|合同价款|价款总额)'
             r'[^0-9¥￥]{0,15}[¥￥]?\s*(\d[\d,，]*\.?\d*)\s*(万元|元)?', True),
            # 人民币 + 数字
            (r'人民币\s*[¥￥]?\s*(\d[\d,，]*\.?\d+)\s*(万元|元)?', True),
            # ¥/￥ + 数字
            (r'[¥￥]\s*(\d[\d,，]*\.?\d+)\s*(万元|元)?', True),
            # 数字 + 万元（宽松兜底）
            (r'(\d[\d,，]{1,8}\.?\d*)\s*(万元)', True),
        ]
        def _parse(raw_num, unit):
            try:
                v = float(raw_num.replace(',', '').replace('，', ''))
                if unit and '万' in unit:
                    v *= 10000
                return f'{v:.2f}'
            except Exception:
                return raw_num

        for line in lines:
            for pat, _ in amount_pats:
                m = re.search(pat, line)
                if m:
                    unit = m.group(2) if m.lastindex and m.lastindex >= 2 else ''
                    result['amount'] = _parse(m.group(1), unit or '')
                    return
        # 全文兜底
        for pat, _ in amount_pats:
            m = re.search(pat, text)
            if m:
                unit = m.group(2) if m.lastindex and m.lastindex >= 2 else ''
                result['amount'] = _parse(m.group(1), unit or '')
                return


def generate_contract_filename(data: dict, original_ext: str) -> str:
    """生成合同新文件名：签订日期_合同名称_甲方关键字_乙方关键字_金额元.ext"""
    date          = data.get('sign_date') or '0000-01-01'
    contract_name = (data.get('contract_name') or '合同')[:15]
    party_a_abbr  = _abbreviate_party(data.get('party_a') or '')
    party_b_abbr  = _abbreviate_party(data.get('party_b') or '')
    amount        = data.get('amount') or ''

    parts = [date, contract_name]
    if party_a_abbr:
        parts.append(party_a_abbr)
    if party_b_abbr:
        parts.append(party_b_abbr)
    if amount:
        try:
            parts.append(f'{float(amount):.0f}元' if float(amount) == int(float(amount))
                         else f'{amount}元')
        except Exception:
            parts.append(f'{amount}元')

    new_name = '_'.join(parts) + original_ext
    new_name = re.sub(r'[\\/:*?"<>|]', '', new_name)
    new_name = re.sub(r'_+', '_', new_name).strip('_')
    return new_name or f"contract_{datetime.now().strftime('%Y%m%d%H%M%S')}{original_ext}"


def _extract_text_from_docx(docx_path: str) -> str:
    """从 Word (.docx / .doc) 文件提取纯文本"""
    # 优先尝试 python-docx（支持 .docx）
    try:
        import docx as _docx
        doc = _docx.Document(docx_path)
        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        if text:
            return text
    except Exception:
        pass

    # 回退：尝试读取 .doc 文件（旧格式，python-docx 不支持）
    # 方法1：尝试 doc2docx 转换
    try:
        from doc2docx import convert as _convert
        import tempfile
        tmp_dir = tempfile.mkdtemp()
        tmp_docx = os.path.join(tmp_dir, 'converted.docx')
        _convert(docx_path, tmp_docx)
        import docx as _docx2
        doc2 = _docx2.Document(tmp_docx)
        text = '\n'.join(p.text for p in doc2.paragraphs if p.text.strip())
        shutil.rmtree(tmp_dir, ignore_errors=True)
        if text:
            return text
    except Exception:
        pass

    # 方法2：尝试用 win32com（仅 Windows 有效）
    if sys.platform == 'win32':
        try:
            import win32com.client as _win32
            word = _win32.Dispatch('Word.Application')
            doc = word.Documents.Open(docx_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            if text and text.strip():
                return text.strip()
        except Exception:
            pass

    # 方法3：尝试作为纯文本读取（某些 .doc 文件可以部分提取）
    try:
        with open(docx_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
            # 过滤掉二进制噪音
            clean = re.sub(r'[^\x20-\x7E\u4e00-\u9fa5\n\r\t]', '', text)
            if len(clean) > 50:
                return clean
    except Exception:
        pass

    return ''


def _pdf_contract_pages_text(pdf_path: str, reader, inv_extractor, max_pages: int = 8) -> str:
    """对合同 PDF 多页做 OCR，合并文本（首页 + 前max_pages页）"""
    try:
        import fitz
        doc = fitz.open(pdf_path)
        combined = []
        for i in range(min(max_pages, len(doc))):
            page = doc[i]
            mat  = fitz.Matrix(2, 2)
            pix  = page.get_pixmap(matrix=mat)
            from PIL import Image
            import io as _io
            img = Image.open(_io.BytesIO(pix.tobytes('png')))
            img_arr = inv_extractor._image_file_to_array.__func__
            # 直接转 numpy
            import numpy as np
            img_np = np.array(img.convert('RGB'))
            img_np = inv_extractor._preprocess_image(img_np)
            results = reader.readtext(img_np, detail=0)
            combined.append('\n'.join(results))
        return '\n'.join(combined)
    except Exception as e:
        return ''


def _image_to_pdf(image_path: str, output_pdf_path: str):
    """将图片文件转为单页 PDF"""
    from PIL import Image as _PIL_Image
    img = _PIL_Image.open(image_path).convert('RGB')
    img.save(output_pdf_path, 'PDF', resolution=150)


_IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif'}


def _resolve_output_ext(orig_ext: str, doc_type: str) -> str:
    """决定输出文件扩展名。
    合同：Word→Word，图片→PDF，PDF→PDF。
    发票/火车票：保持原格式不变。"""
    if doc_type == 'contract':
        if orig_ext in ('.docx', '.doc'):
            return orig_ext          # Word 合同保持 Word
        if orig_ext in _IMAGE_EXTS:
            return '.pdf'            # 图片合同转为 PDF
    return orig_ext                  # 其他格式原样


def _generate_any_filename(data: dict, doc_type: str, out_ext: str) -> str:
    """根据文档类型选择对应的命名函数"""
    if doc_type == 'train':
        return generate_train_filename(data, out_ext)
    if doc_type == 't3':
        return generate_t3_filename(data, out_ext)
    if doc_type == 'contract':
        return generate_contract_filename(data, out_ext)
    return generate_filename(data, out_ext)


def _save_file(src_path: str, src_ext: str, dst_path: str,
               out_ext: str, doc_type: str):
    """将源文件保存到 dst_path，必要时做格式转换（图片合同→PDF）。"""
    if doc_type == 'contract' and src_ext in _IMAGE_EXTS and out_ext == '.pdf':
        _image_to_pdf(src_path, dst_path)
    else:
        shutil.copy(src_path, dst_path)


# ======================== Flask 应用 ========================

app = Flask(__name__, template_folder='templates')
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024

READER = None
READER_READY = False
UPLOAD_RESULTS = {}

# ---- 全局批次 (单服务器单用户) ----
BATCH_RESULTS = []        # 累积的所有 result 条目
BATCH_SESSION_DIRS = []   # 累积的 session 目录（用于下载 ZIP）
BATCH_LOCK = threading.Lock()


def init_ocr_background():
    """后台初始化 OCR"""
    global READER, READER_READY
    try:
        print("正在初始化 OCR 模型...")
        READER = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)
        READER_READY = True
        print("✅ OCR 模型已准备好")
    except Exception as e:
        print(f"❌ OCR 初始化失败: {e}")
        READER_READY = False


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/status')
def status():
    """获取 OCR 准备状态"""
    return jsonify({
        'ready': READER_READY,
        'message': 'OCR 已准备好' if READER_READY else '正在初始化 OCR 模型，请稍候...'
    })


def _pdf_direct_text(pdf_path: str) -> str | None:
    """用 PyMuPDF 直接提取矢量 PDF 文本（电子发票/电子客票专用）。
    若提取到足够中文字符（>=10），返回按行合并的文本；否则返回 None 让调用方退回 OCR。"""
    try:
        import fitz
        doc = fitz.open(pdf_path)
        lines = []
        for page in doc:
            for line in page.get_text().splitlines():
                line = line.strip()
                if line:
                    lines.append(line)
        doc.close()
        full = '\n'.join(lines)
        cn_chars = sum(1 for c in full if '\u4e00' <= c <= '\u9fa5')
        return full if cn_chars >= 10 else None
    except Exception:
        return None


def smart_extract(file_path: str, reader) -> tuple:
    """自动判断发票/火车票/合同，返回 (data_dict, doc_type)。
    PDF 优先用 PyMuPDF 直提矢量文本；Word 用 python-docx；图片 OCR 兜底。"""
    inv      = InvoiceExtractor(reader)
    train    = TrainTicketExtractor(reader)
    contract = ContractExtractor()

    fp  = Path(file_path)
    ext = fp.suffix.lower()

    text = None  # 最终用于字段提取的文本

    # === 1. Word 文件：直接提取文本 ===
    if ext in ('.docx', '.doc'):
        text = _extract_text_from_docx(str(fp))
        if not text:
            raise Exception("无法读取 Word 文件内容")

    # === 2. 矢量 PDF 优先直提文本 ===
    elif ext == '.pdf':
        text = _pdf_direct_text(str(fp))

    # === 3. 无法直提 / 图片 → 单页 OCR（第1页预判类型）===
    if not text:
        if ext == '.pdf':
            img = inv._pdf_to_image(str(fp))
        else:
            img = inv._image_file_to_array(str(fp))

        if img is None or img.size == 0:
            raise Exception("无法读取图像")

        img     = inv._preprocess_image(img)
        results = reader.readtext(img, detail=0)
        if not results:
            raise Exception("OCR 无法识别")
        page1_text = '\n'.join(results)

        # 若第1页已判断为合同，追加后续多页 OCR（合同关键信息常在后页）
        if detect_doc_type(page1_text) == 'contract' and ext == '.pdf':
            extra = _pdf_contract_pages_text(str(fp), reader, inv, max_pages=8)
            text = page1_text + '\n' + extra if extra else page1_text
        else:
            text = page1_text

    # 若直提的文本判断为合同，补充多页 OCR（矢量PDF合同文本可能不完整）
    if text and detect_doc_type(text) == 'contract' and ext == '.pdf':
        extra = _pdf_contract_pages_text(str(fp), reader, inv, max_pages=8)
        if extra:
            text = text + '\n' + extra

    doc_type = detect_doc_type(text)

    if doc_type == 'train':
        fields = train._extract_fields(text, fp.stem)
    elif doc_type == 't3':
        # T3出行/网约车：使用火车票提取逻辑（需要出发地/到达地），但命名用T3格式
        fields = train._extract_fields(text, fp.stem)
        fields['invoice_type'] = 'T3出行'
    elif doc_type == 'contract':
        fields = contract._extract_fields(text)
    else:
        fields = inv._extract_fields(text)
        # 从文件名补充发票缺失字段
        fn_fields = inv._extract_from_filename(fp.stem)
        for key in ('date', 'invoice_number', 'buyer', 'supplier'):
            if not fields.get(key) and fn_fields.get(key):
                fields[key] = fn_fields[key]

    fields['_raw_text'] = text
    return fields, doc_type


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """处理上传 - 保存文件，识别，重命名"""
    try:
        if 'files[]' not in request.files:
            return jsonify({'error': '未选择文件'}), 400

        if not READER_READY:
            return jsonify({'error': 'OCR 还未初始化，请稍候片刻后重试'}), 503

        files = request.files.getlist('files[]')
        reader = READER
        if reader is None:
            return jsonify({'error': 'OCR 初始化失败'}), 500

        results = []

        session_id  = datetime.now().strftime('%Y%m%d%H%M%S%f')[-15:]
        session_dir = os.path.join(tempfile.gettempdir(), f'invoice_session_{session_id}')
        os.makedirs(session_dir, exist_ok=True)

        for file in files:
            if file.filename == '':
                continue

            ext = Path(file.filename).suffix.lower()
            allowed_ext = {
                '.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif',
                '.docx', '.doc', '.zip'
            }

            if ext not in allowed_ext:
                results.append({'filename': file.filename, 'status': 'error', 'error': '不支持的格式'})
                continue

            safe_name = re.sub(r'[^\w.\-]', '_', file.filename)
            temp_path = os.path.join(tempfile.gettempdir(), f"{session_id}_{safe_name}")
            file.save(temp_path)

            try:
                if ext == '.zip':
                    zip_results = process_zip_file(temp_path, reader, session_dir)
                    results.extend(zip_results)
                else:
                    data, doc_type = smart_extract(temp_path, reader)
                    out_ext = _resolve_output_ext(ext, doc_type)
                    new_name = _generate_any_filename(data, doc_type, out_ext)

                    new_path = os.path.join(session_dir, new_name)
                    _save_file(temp_path, ext, new_path, out_ext, doc_type)

                    results.append({
                        'filename': file.filename,
                        'new_name': new_name,
                        'data': data,
                        'doc_type': doc_type,
                        'status': 'success'
                    })
            except Exception as e:
                results.append({'filename': file.filename, 'status': 'error', 'error': str(e)})
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        UPLOAD_RESULTS[session_id] = {'results': results, 'session_dir': session_dir}
        return jsonify({'session_id': session_id, 'total': len(results), 'results': results})

    except Exception as e:
        return jsonify({'error': f'处理失败: {str(e)}'}), 500


def process_zip_file(zip_path: str, reader, session_dir: str) -> list:
    """处理 ZIP 文件，自动识别发票/火车票/合同"""
    results      = []
    temp_extract = tempfile.mkdtemp()

    try:
        with zipfile.ZipFile(zip_path, 'r') as z:
            z.extractall(temp_extract)

        allowed_ext = {
            '.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif',
            '.docx', '.doc'
        }

        for root, dirs, files in os.walk(temp_extract):
            for file in sorted(files):
                ext = Path(file).suffix.lower()
                if ext in allowed_ext:
                    file_path = os.path.join(root, file)
                    try:
                        data, doc_type = smart_extract(file_path, reader)
                        out_ext  = _resolve_output_ext(ext, doc_type)
                        new_name = _generate_any_filename(data, doc_type, out_ext)
                        new_path = os.path.join(session_dir, new_name)
                        _save_file(file_path, ext, new_path, out_ext, doc_type)

                        results.append({
                            'filename': file,
                            'new_name': new_name,
                            'data': data,
                            'doc_type': doc_type,
                            'status': 'success'
                        })
                    except Exception as e:
                        results.append({'filename': file, 'status': 'error', 'error': str(e)})
    finally:
        shutil.rmtree(temp_extract, ignore_errors=True)

    return results


@app.route('/api/download-zip/<session_id>', methods=['GET'])
def download_zip(session_id):
    """将本次会话所有重命名文件打包成 ZIP 并下载（浏览器弹出另存为对话框）"""
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 400
    session_data = UPLOAD_RESULTS[session_id]
    results  = session_data['results']
    sdir     = session_data['session_dir']
    success  = [item for item in results if item['status'] == 'success']
    if not success:
        return jsonify({'error': '无可下载文件'}), 400
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for item in success:
            fpath = os.path.join(sdir, item['new_name'])
            if os.path.exists(fpath):
                zf.write(fpath, item['new_name'])
    zip_buf.seek(0)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return send_file(zip_buf, mimetype='application/zip',
                     as_attachment=True,
                     download_name=f'发票_{ts}.zip')


@app.route('/api/download/<session_id>', methods=['GET'])
def download_results(session_id):
    """返回本次会话所有可下载文件的列表（文件名 + 下载 URL）"""
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 400
    results = UPLOAD_RESULTS[session_id]['results']
    files = [
        {'new_name': item['new_name'], 'index': i}
        for i, item in enumerate(results)
        if item['status'] == 'success'
    ]
    return jsonify({'files': files})


@app.route('/api/download-file/<session_id>/<int:item_idx>', methods=['GET'])
def download_single_file(session_id, item_idx):
    """下载单个重命名后的文件"""
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 400
    session_data = UPLOAD_RESULTS[session_id]
    results = session_data['results']
    if item_idx < 0 or item_idx >= len(results):
        return jsonify({'error': '索引越界'}), 400
    item = results[item_idx]
    if item['status'] != 'success':
        return jsonify({'error': '该文件未成功识别'}), 400
    new_name = item['new_name']
    file_path = os.path.join(session_data['session_dir'], new_name)
    if not os.path.exists(file_path):
        return jsonify({'error': '文件已过期'}), 400
    return send_file(file_path, as_attachment=True, download_name=new_name)


def _build_csv_bytes(results: list) -> bytes:
    """将 result 列表序列化为带 BOM 的 UTF-8 CSV bytes（支持发票 + 火车票 + 合同）"""
    import csv
    output = io.StringIO()
    output.write('\ufeff')
    writer = csv.writer(output)
    writer.writerow([
        '原文件名', '新文件名', '状态', '类型', '日期',
        # 发票专属
        '发票号码', '购买方', '销售方', '金额（不含税）', '税额', '价税合计',
        # 火车票专属
        '车次', '出发站', '到达站', '乘客姓名', '座位', '座位类型', '票价',
        # 合同专属
        '合同名称', '甲方', '乙方', '合同金额',
        '错误信息'
    ])

    total_invoice = total_train = total_contract = 0.0
    success_count = fail_count = 0
    EMPTY6 = [''] * 6
    EMPTY7 = [''] * 7
    EMPTY4 = [''] * 4

    for item in results:
        if item['status'] == 'success':
            d     = item.get('data') or {}
            dtype = item.get('doc_type', 'invoice')

            if dtype == 'train':
                price = d.get('price', '')
                writer.writerow([
                    item.get('filename', ''), item.get('new_name', ''), '成功', '🚄 火车票',
                    d.get('date', ''),
                    *EMPTY6,                          # 发票列留空
                    d.get('train_number', ''), d.get('from_station', ''),
                    d.get('to_station', ''), d.get('passenger_name', ''),
                    d.get('seat', ''), d.get('seat_type', ''), price,
                    *EMPTY4, ''                        # 合同列留空 + 错误
                ])
                try: total_train += float(price) if price else 0
                except ValueError: pass

            elif dtype == 't3':
                # T3出行/网约车：使用火车票字段但显示为出行服务
                price = d.get('price') or d.get('amount', '')
                writer.writerow([
                    item.get('filename', ''), item.get('new_name', ''), '成功', '🚗 T3出行',
                    d.get('date', ''),
                    *EMPTY6,                          # 发票列留空
                    d.get('train_number', ''), d.get('from_station', ''),
                    d.get('to_station', ''), d.get('passenger_name', ''),
                    d.get('seat', ''), d.get('seat_type', ''), price,
                    *EMPTY4, ''                        # 合同列留空 + 错误
                ])
                try: total_train += float(price) if price else 0
                except ValueError: pass

            elif dtype == 'contract':
                amount = d.get('amount', '')
                writer.writerow([
                    item.get('filename', ''), item.get('new_name', ''), '成功', '📋 合同',
                    d.get('sign_date', ''),
                    *EMPTY6,                          # 发票列留空
                    *EMPTY7,                          # 火车票列留空
                    d.get('contract_name', ''), d.get('party_a', ''),
                    d.get('party_b', ''), amount, ''
                ])
                try: total_contract += float(amount) if amount else 0
                except ValueError: pass

            else:
                tax_free = d.get('tax_free_amount', '')
                tax      = d.get('tax_amount', '')
                amount   = d.get('amount', '')
                writer.writerow([
                    item.get('filename', ''), item.get('new_name', ''), '成功', '🧾 发票',
                    d.get('date', ''),
                    d.get('invoice_number', ''), d.get('buyer', ''), d.get('supplier', ''),
                    tax_free, tax, amount,
                    *EMPTY7,                          # 火车票列留空
                    *EMPTY4, ''                        # 合同列留空 + 错误
                ])
                try: total_invoice += float(amount) if amount else 0
                except ValueError: pass

            success_count += 1
        else:
            writer.writerow([
                item.get('filename', ''), '', '失败', '', '',
                *EMPTY6, *EMPTY7, *EMPTY4,
                item.get('error', '')
            ])
            fail_count += 1

    writer.writerow([])
    writer.writerow([
        f'合计（成功 {success_count} 张，失败 {fail_count} 张）',
        '', '', '', '',
        '', '', '', '', '', f'{total_invoice:.2f}',
        '', '', '', '', '', '', f'{total_train:.2f}',
        '', '', '', f'{total_contract:.2f}', ''
    ])
    return output.getvalue().encode('utf-8-sig')


def _build_excel_bytes(results: list, title: str = '识别结果汇总') -> bytes:
    """将 result 列表序列化为带格式的 .xlsx bytes（支持发票 + 火车票 + 合同）"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '汇总表'

    # ---- 颜色 & 样式常量 ----
    C_INVOICE_HDR  = 'FF276749'   # 深绿（发票表头）
    C_TRAIN_HDR    = 'FF2B6CB0'   # 深蓝（火车票表头）
    C_CONTRACT_HDR = 'FF6B21A8'   # 深紫（合同表头）
    C_META_HDR     = 'FF4A5568'   # 深灰（公共表头）
    C_INVOICE_ROW  = 'FFE6F4EA'   # 浅绿（发票行底色）
    C_TRAIN_ROW    = 'FFE8F0FE'   # 浅蓝（火车票行底色）
    C_CONTRACT_ROW = 'FFF3E8FE'   # 浅紫（合同行底色）
    C_FAIL_ROW     = 'FFFFF3CD'   # 浅黄（失败行底色）
    C_SUM_ROW      = 'FFFFF8E1'   # 浅橙（汇总行底色）
    WHITE          = 'FFFFFFFF'

    thin = Side(style='thin', color='FFCCCCCC')
    bdr  = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hdr_font(color='FFFFFFFF', bold=True):
        return Font(name='微软雅黑', bold=bold, color=color, size=10)

    def cell_font(bold=False, color='FF333333'):
        return Font(name='微软雅黑', bold=bold, color=color, size=9)

    def fill(hex_color):
        return PatternFill('solid', fgColor=hex_color)

    def center():
        return Alignment(horizontal='center', vertical='center', wrap_text=True)

    def left():
        return Alignment(horizontal='left', vertical='center', wrap_text=True)

    # ---- 表头定义（列索引 1-based）----
    # 公共(1-5) | 发票(6-11) | 火车票(12-18) | 合同(19-22) | 错误(23)
    HEADERS = [
        '原文件名', '新文件名', '状态', '类型', '日期',
        '发票号码', '购买方', '销售方', '金额(不含税)', '税额', '价税合计',
        '车次', '出发站', '到达站', '乘客姓名', '座位', '座位类型', '票价',
        '合同名称', '甲方', '乙方', '合同金额',
        '错误信息',
    ]
    NCOLS = len(HEADERS)  # 23
    COL_WIDTHS = [26, 36, 8, 10, 12,
                  18, 20, 20, 14, 12, 12,
                  10, 14, 14, 14, 14, 10, 10,
                  20, 16, 16, 14,
                  20]
    HDR_COLORS = {
        **{i: C_META_HDR     for i in range(1, 6)},
        **{i: C_INVOICE_HDR  for i in range(6, 12)},
        **{i: C_TRAIN_HDR    for i in range(12, 19)},
        **{i: C_CONTRACT_HDR for i in range(19, 23)},
        23: C_META_HDR,
    }

    # ---- 标题行 ----
    ws.merge_cells(f'A1:{get_column_letter(NCOLS)}1')
    title_cell = ws['A1']
    title_cell.value     = title
    title_cell.font      = Font(name='微软雅黑', bold=True, size=14, color=WHITE)
    title_cell.fill      = fill('FF4A5568')
    title_cell.alignment = center()
    ws.row_dimensions[1].height = 32

    for col_i, (hdr, width) in enumerate(zip(HEADERS, COL_WIDTHS), start=1):
        c = ws.cell(row=2, column=col_i, value=hdr)
        c.font      = hdr_font(color=WHITE)
        c.fill      = fill(HDR_COLORS.get(col_i, C_META_HDR))
        c.alignment = center()
        c.border    = bdr
        ws.column_dimensions[get_column_letter(col_i)].width = width
    ws.row_dimensions[2].height = 22

    # ---- 数据行 ----
    total_invoice = total_train = total_contract = 0.0
    success_count = fail_count  = 0
    data_start_row = 3

    for row_i, item in enumerate(results, start=data_start_row):
        dtype    = item.get('doc_type', 'invoice')
        is_ok    = item['status'] == 'success'
        d        = item.get('data') or {}

        # 行底色
        if not is_ok:
            row_fill = fill(C_FAIL_ROW)
        elif dtype == 'train':
            row_fill = fill(C_TRAIN_ROW)
        elif dtype == 't3':
            row_fill = fill(C_TRAIN_ROW)  # T3出行用火车票底色
        elif dtype == 'contract':
            row_fill = fill(C_CONTRACT_ROW)
        else:
            row_fill = fill(C_INVOICE_ROW)

        def set_cell(col, val, bold=False, num_fmt=None, align=None):
            c = ws.cell(row=row_i, column=col, value=val)
            c.font      = cell_font(bold=bold)
            c.fill      = row_fill
            c.border    = bdr
            c.alignment = align or left()
            if num_fmt:
                c.number_format = num_fmt

        def num_cell(col, val_str):
            try:
                v = float(val_str) if val_str else None
                c = ws.cell(row=row_i, column=col, value=v)
                c.font = cell_font(); c.fill = row_fill; c.border = bdr
                c.alignment = center(); c.number_format = '#,##0.00'
                return v or 0
            except (ValueError, TypeError):
                set_cell(col, val_str, align=center())
                return 0

        if is_ok:
            # 公共列
            type_labels = {'train': '火车票', 't3': '🚗 T3出行', 'contract': '合同', 'invoice': '发票'}
            set_cell(1, item.get('filename', ''))
            set_cell(2, item.get('new_name', ''), bold=True)
            set_cell(3, '成功', align=center())
            set_cell(4, type_labels.get(dtype, '发票'), align=center())

            if dtype in ('train', 't3'):
                set_cell(5, d.get('date', ''), align=center())
                for col in range(6, 12): set_cell(col, '')
                set_cell(12, d.get('train_number', ''), align=center())
                set_cell(13, d.get('from_station', ''))
                set_cell(14, d.get('to_station', ''))
                set_cell(15, d.get('passenger_name', ''))
                set_cell(16, d.get('seat', ''), align=center())
                set_cell(17, d.get('seat_type', ''), align=center())
                price_str = d.get('price', '')
                total_train += num_cell(18, price_str)
                for col in range(19, 23): set_cell(col, '')

            elif dtype == 'contract':
                set_cell(5, d.get('sign_date', ''), align=center())
                for col in range(6, 19): set_cell(col, '')
                set_cell(19, d.get('contract_name', ''))
                set_cell(20, d.get('party_a', ''))
                set_cell(21, d.get('party_b', ''))
                total_contract += num_cell(22, d.get('amount', ''))

            else:  # invoice
                set_cell(5, d.get('date', ''), align=center())
                set_cell(6,  d.get('invoice_number', ''), align=center())
                set_cell(7,  d.get('buyer', ''))
                set_cell(8,  d.get('supplier', ''))
                num_cell(9,  d.get('tax_free_amount', ''))
                num_cell(10, d.get('tax_amount', ''))
                total_invoice += num_cell(11, d.get('amount', ''))
                for col in range(12, 23): set_cell(col, '')

            set_cell(23, '')
            success_count += 1
        else:
            set_cell(1, item.get('filename', ''))
            for col in range(2, 23): set_cell(col, '')
            set_cell(3, '失败', align=center())
            set_cell(23, item.get('error', ''))
            fail_count += 1

        ws.row_dimensions[row_i].height = 18

    # ---- 汇总行 ----
    sum_row = data_start_row + len(results)
    ws.merge_cells(f'A{sum_row}:D{sum_row}')
    sc = ws.cell(row=sum_row, column=1,
                 value=f'合计：成功 {success_count} 张，失败 {fail_count} 张')
    sc.font = Font(name='微软雅黑', bold=True, size=10, color='FF333333')
    sc.fill = fill(C_SUM_ROW); sc.alignment = center(); sc.border = bdr

    def sum_cell(col, val, color):
        c = ws.cell(row=sum_row, column=col, value=val)
        c.font = Font(name='微软雅黑', bold=True, size=10, color=color)
        c.fill = fill(C_SUM_ROW); c.border = bdr
        c.alignment = center(); c.number_format = '#,##0.00'

    def sum_lbl(col, txt):
        c = ws.cell(row=sum_row, column=col, value=txt)
        c.font = Font(name='微软雅黑', bold=True, size=9)
        c.fill = fill(C_SUM_ROW); c.border = bdr; c.alignment = center()

    sum_lbl(10, '发票总额 ▶'); sum_cell(11, total_invoice, C_INVOICE_HDR[2:])
    sum_lbl(17, '火车票总额 ▶'); sum_cell(18, total_train,   C_TRAIN_HDR[2:])
    sum_lbl(21, '合同总额 ▶');  sum_cell(22, total_contract, C_CONTRACT_HDR[2:])

    for col in [5,6,7,8,9,12,13,14,15,16,19,20,23]:
        c = ws.cell(row=sum_row, column=col, value='')
        c.fill = fill(C_SUM_ROW); c.border = bdr

    ws.row_dimensions[sum_row].height = 22
    ws.freeze_panes = 'A3'

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


@app.route('/api/export/<session_id>', methods=['GET'])
def export_csv(session_id):
    """导出本次识别结果为 CSV 表格"""
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 400
    csv_bytes = _build_csv_bytes(UPLOAD_RESULTS[session_id]['results'])
    return send_file(
        io.BytesIO(csv_bytes),
        mimetype='text/csv; charset=utf-8-sig',
        as_attachment=True,
        download_name=f'发票识别结果_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
    )


@app.route('/api/export-excel/<session_id>', methods=['GET'])
def export_excel(session_id):
    """导出本次识别结果为 Excel 表格"""
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 400
    results = UPLOAD_RESULTS[session_id]['results']
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    xlsx_bytes = _build_excel_bytes(results, title=f'识别结果汇总 — {ts}')
    return send_file(
        io.BytesIO(xlsx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'发票识别结果_{ts}.xlsx'
    )


# ======================== 批次管理 API ========================

@app.route('/api/update/<session_id>/<int:item_index>', methods=['POST'])
def update_item(session_id, item_index):
    """手动修改识别字段（发票或火车票），并同步重命名实体文件"""
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 404

    results     = UPLOAD_RESULTS[session_id]['results']
    session_dir = UPLOAD_RESULTS[session_id]['session_dir']

    if item_index < 0 or item_index >= len(results):
        return jsonify({'error': '索引越界'}), 400

    item = results[item_index]
    if item['status'] != 'success':
        return jsonify({'error': '只能编辑成功识别的条目'}), 400

    payload      = request.get_json(force=True, silent=True) or {}
    old_new_name = item.get('new_name', '')
    doc_type     = item.get('doc_type', 'invoice')
    d            = item.get('data') or {}

    # 更新字段（发票 + 火车票 + 合同字段均接受，只更新有值的）
    all_fields = ('date', 'invoice_number', 'buyer', 'supplier',
                  'tax_free_amount', 'tax_amount', 'amount',
                  'train_number', 'depart_time', 'from_station', 'to_station',
                  'passenger_name', 'seat', 'seat_type', 'price',
                  'contract_name', 'sign_date', 'party_a', 'party_b')
    for field in all_fields:
        if field in payload:
            d[field] = payload[field].strip()
    item['data'] = d

    ext          = os.path.splitext(old_new_name)[1].lower()
    new_new_name = _generate_any_filename(d, doc_type, ext)
    item['new_name'] = new_new_name

    old_path = os.path.join(session_dir, old_new_name)
    new_path = os.path.join(session_dir, new_new_name)
    if os.path.exists(old_path) and old_path != new_path:
        try:
            os.rename(old_path, new_path)
        except OSError:
            pass

    results[item_index] = item
    return jsonify({'new_name': new_new_name, 'data': d, 'doc_type': doc_type})


@app.route('/api/batch/add/<session_id>', methods=['POST'])
def batch_add(session_id):
    """把一次会话的结果追加到全局批次"""
    global BATCH_RESULTS, BATCH_SESSION_DIRS
    if session_id not in UPLOAD_RESULTS:
        return jsonify({'error': '会话已过期'}), 404
    session_data = UPLOAD_RESULTS[session_id]
    with BATCH_LOCK:
        for idx, item in enumerate(session_data['results']):
            # 记录来源 session 和索引，供批次逐文件下载使用
            item_copy = dict(item)
            item_copy['session_id'] = session_id
            item_copy['_idx']       = idx
            BATCH_RESULTS.append(item_copy)
        BATCH_SESSION_DIRS.append(session_data['session_dir'])
        total   = len(BATCH_RESULTS)
        success = sum(1 for r in BATCH_RESULTS if r['status'] == 'success')
    return jsonify({'total': total, 'success': success})


@app.route('/api/batch/status', methods=['GET'])
def batch_status():
    """返回当前批次累积统计"""
    with BATCH_LOCK:
        total   = len(BATCH_RESULTS)
        success = sum(1 for r in BATCH_RESULTS if r['status'] == 'success')
    return jsonify({'total': total, 'success': success})


@app.route('/api/batch/files', methods=['GET'])
def batch_files():
    """返回批次中所有可下载文件的列表（供前端逐文件写入文件夹用）"""
    with BATCH_LOCK:
        results      = list(BATCH_RESULTS)
        session_dirs = list(BATCH_SESSION_DIRS)
    if not results:
        return jsonify({'files': []})

    files = []
    # 按照各 session_dir 逐文件扫描，与 BATCH_RESULTS 顺序一致
    seen_names: dict = {}
    for sdir in session_dirs:
        if not os.path.isdir(sdir):
            continue
        for item in results:
            fname = item.get('new_name', '')
            if not fname:
                continue
            fpath = os.path.join(sdir, fname)
            if not os.path.exists(fpath):
                continue
            # 取 session_id 供前端复用 download-file 接口
            sid = item.get('session_id', '')
            idx = item.get('_idx', None)
            if sid and idx is not None:
                files.append({'new_name': fname, 'session_id': sid, 'index': idx})
    return jsonify({'files': files})


@app.route('/api/batch/download', methods=['GET'])
def batch_download():
    """下载批次中所有重命名文件的 ZIP"""
    with BATCH_LOCK:
        session_dirs = list(BATCH_SESSION_DIRS)
        results      = list(BATCH_RESULTS)
    if not results:
        return jsonify({'error': '批次为空'}), 404

    zip_buf = io.BytesIO()
    seen: dict = {}
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for sdir in session_dirs:
            if not os.path.isdir(sdir):
                continue
            for fname in os.listdir(sdir):
                fpath = os.path.join(sdir, fname)
                arcname = fname
                if arcname in seen:
                    seen[arcname] += 1
                    stem, ext = os.path.splitext(fname)
                    arcname = f'{stem}_{seen[fname]}{ext}'
                else:
                    seen[arcname] = 1
                zf.write(fpath, arcname)
    zip_buf.seek(0)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return send_file(zip_buf, mimetype='application/zip',
                     as_attachment=True,
                     download_name=f'发票批次_{ts}.zip')


@app.route('/api/batch/export', methods=['GET'])
def batch_export():
    """导出批次全部结果为 CSV"""
    with BATCH_LOCK:
        results = list(BATCH_RESULTS)
    if not results:
        return jsonify({'error': '批次为空'}), 404
    csv_bytes = _build_csv_bytes(results)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return send_file(io.BytesIO(csv_bytes),
                     mimetype='text/csv; charset=utf-8-sig',
                     as_attachment=True,
                     download_name=f'发票批次汇总_{ts}.csv')


@app.route('/api/batch/export-excel', methods=['GET'])
def batch_export_excel():
    """导出批次全部结果为 Excel"""
    with BATCH_LOCK:
        results = list(BATCH_RESULTS)
    if not results:
        return jsonify({'error': '批次为空'}), 404
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    xlsx_bytes = _build_excel_bytes(results, title=f'批次汇总表 — {ts}')
    return send_file(
        io.BytesIO(xlsx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'发票批次汇总_{ts}.xlsx'
    )


@app.route('/api/batch/clear', methods=['POST'])
def batch_clear():
    """清空全局批次"""
    global BATCH_RESULTS, BATCH_SESSION_DIRS
    with BATCH_LOCK:
        BATCH_RESULTS = []
        BATCH_SESSION_DIRS = []
    return jsonify({'message': '批次已清空'})


@app.errorhandler(413)
def file_too_large(e):
    return jsonify({'error': '文件过大，最大支持 500MB'}), 413


@app.errorhandler(Exception)
def handle_exception(e):
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException):
        return e
    import traceback
    print(f"未处理异常: {traceback.format_exc()}")
    return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500


@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': '服务器内部错误，请重试'}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') != 'production'

    print(f"\n{'='*60}")
    print("发票批量重命名工具")
    print(f"{'='*60}")
    print(f"访问地址: http://127.0.0.1:{port}")
    print(f"{'='*60}\n")

    ocr_thread = threading.Thread(target=init_ocr_background, daemon=True)
    ocr_thread.start()

    try:
        app.run(host='0.0.0.0', port=port, debug=debug, threaded=True, use_reloader=False)
    except OSError as e:
        if "Address already in use" in str(e):
            print(f"\n❌ 错误: 端口 {port} 已被占用")
        else:
            print(f"\n❌ 启动失败: {e}")
    except KeyboardInterrupt:
        print("\n已停止服务")